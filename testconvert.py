import fitz  # PyMuPDF for pdf conversion
from docx import Document # for Microsoft word file conversion
import mammoth 
from bs4 import BeautifulSoup
import re


def pdf_to_text(pdf_path):
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def docx_to_text_old(docx_path):
    doc = Document(docx_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return text

def extract_docx_header_footer(docx_path):
    # Load the document using python-docx
    doc = Document(docx_path)

    # Extract header text
    header_text = []
    for section in doc.sections:
        for header in section.header.paragraphs:
            header_text.append(header.text)

    # Extract footer text
    footer_text = []
    for section in doc.sections:
        for footer in section.footer.paragraphs:
            footer_text.append(footer.text)
    
    result = "\n".join(header_text + footer_text)
    result = re.sub('\n+','\n', result)
    return result

def docx_to_text(docx_path):
    # Read and convert temp Word file into HTML for analysis
    with open(docx_path, "rb") as docx_file:
        try:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
        except Exception as e:
            html = f"<p>docx_to_text ERROR: Document was not a Microsoft Word document in .docx format (Zip file) that could be analyzed. Exception {e}</p>\n"

    # Get the header and footer from the file
    headerfooter = extract_docx_header_footer(docx_path)
    # Create a BeautifulSoup object and extract all the text
    soup = BeautifulSoup(html, 'html.parser')
    plain_text = soup.get_text(separator='\n')
    #plain_text = "HEADER_FOOTER: " + headerfooter + '\nFRONT_PAGE:\n' + plain_text
    plain_text = headerfooter + '\n' + plain_text
    plain_text = re.sub('\t',' ',plain_text)
    return plain_text

# print(pdf_to_text(".\\Advisor.pdf"))
print(docx_to_text("Testdoc.docx"))